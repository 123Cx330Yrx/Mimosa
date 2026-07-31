from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "SYSTEM_EXPERIMENT_HANDOFF.zh-CN.md"
OUTPUT = Path(r"C:\Users\33502\Documents\Codex\Mimosa\deliverables\Mimosa-System-Experiment-Handoff-Bilingual-2026-07-31.docx")

# compact_reference_guide preset + editorial_cover header pattern
FONT = "Microsoft YaHei"
INK = "173E2E"
ACCENT = "2F7D5C"
MUTED = "64766D"
PALE = "E8F1EC"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(.492)
section.footer_distance = Inches(.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, before, after, color in [
    ("Heading 1", 16, 18, 10, ACCENT),
    ("Heading 2", 13, 14, 7, ACCENT),
    ("Heading 3", 12, 10, 5, INK),
]:
    style = styles[name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for name in ("List Bullet", "List Number"):
    style = styles[name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(11)
    style.paragraph_format.left_indent = Inches(.375)
    style.paragraph_format.first_line_indent = Inches(-.188)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.25

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_run_font(header.add_run("MIMOSA · SYSTEM & STUDY HANDOFF"), size=8.5, color=MUTED, bold=True)
add_page_number(section.footer.paragraphs[0])

# Editorial-cover opening: generous but compact enough for an operational guide.
for _ in range(3):
    doc.add_paragraph()
kicker = doc.add_paragraph()
kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
kicker.paragraph_format.space_after = Pt(14)
set_run_font(kicker.add_run("SRTP RESEARCH PROTOTYPE"), size=10, color=ACCENT, bold=True)
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(8)
set_run_font(title.add_run("Mimosa 系统设计与实验交接手册"), size=28, color=INK, bold=True)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(28)
set_run_font(subtitle.add_run("中英文运行版本 · System Design · Study Operations · Deployment"), size=13, color=MUTED)
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_after = Pt(90)
set_run_font(meta.add_run("协议 v14-bilingual  |  2026-07-31  |  mimosa-srtp.com"), size=10, color=ACCENT, bold=True)
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(note.add_run("面向研究团队的系统、实验与交付参考"), size=10, color=MUTED, italic=True)
doc.add_page_break()

lines = SOURCE.read_text(encoding="utf-8").splitlines()
i = 4
while i < len(lines):
    line = lines[i].rstrip()
    if not line:
        i += 1
        continue
    if line.startswith("## "):
        doc.add_heading(line[3:], level=1)
    elif line.startswith("### "):
        doc.add_heading(line[4:], level=2)
    elif line.startswith("|") and i + 1 < len(lines) and re.fullmatch(r"[| :\-]+", lines[i + 1]):
        headers = [x.strip() for x in line.strip("|").split("|")]
        rows = []
        i += 2
        while i < len(lines) and lines[i].startswith("|"):
            rows.append([x.strip().replace("`", "") for x in lines[i].strip("|").split("|")])
            i += 1
        i -= 1
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        base = TABLE_WIDTH // len(headers)
        widths = [base] * len(headers)
        widths[-1] += TABLE_WIDTH - sum(widths)
        for index, value in enumerate(headers):
            cell = table.rows[0].cells[index]
            cell.text = value
            set_cell_shading(cell, PALE)
            for run in cell.paragraphs[0].runs:
                set_run_font(run, size=9.5, color=INK, bold=True)
        for row in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row[:len(headers)]):
                cells[index].text = value
                for run in cells[index].paragraphs[0].runs:
                    set_run_font(run, size=9.3, color=INK)
        set_table_geometry(table, widths)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
    elif re.match(r"^\d+\.\s", line):
        doc.add_paragraph(re.sub(r"^\d+\.\s", "", line), style="List Number")
    elif line.startswith("- "):
        doc.add_paragraph(line[2:], style="List Bullet")
    else:
        paragraph = doc.add_paragraph()
        text = line.replace("**", "").replace("`", "")
        set_run_font(paragraph.add_run(text), size=11, color=INK)
    i += 1

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.core_properties.title = "Mimosa 系统设计与实验交接手册"
doc.core_properties.subject = "Bilingual system design, study operations, data logging and deployment"
doc.core_properties.author = "Mimosa SRTP Team"
doc.save(OUTPUT)
print(OUTPUT)
